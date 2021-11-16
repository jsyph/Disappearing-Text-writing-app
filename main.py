import datetime
import pathlib
import tkinter as tk
import tkinter.messagebox as tk_message
from tkinter import filedialog
from tkinter import ttk

import docx

from globals import Globals
from init import init


def reset_timer():
    Globals.counter = 10
    Globals.last_press = [datetime.datetime.now()]


def create_new_doc():
    response = tk_message.askokcancel("WARNING", "Do you want to reset the document?")
    if response:
        main_entry.delete('1.0', tk.END)
        reset_timer()


def save_doc():
    def insert_newlines(string, every=100):
        lines = []
        for i in range(0, len(string), every):
            lines.append(string[i:i + every])
        return '\n'.join(lines)

    print("SAVE THE DOCUMENT")
    text = insert_newlines(main_entry.get('1.0', 'end'))

    response = tk_message.askyesno("CHOOSE", "Press 'Yes' if you want to save the file in the current directory."
                                             "\nPress 'No' to select a save file location")

    if not response:
        save_loc = filedialog.asksaveasfilename(initialdir="/", title="Select file",
                                                filetypes=(("Text File", "*.txt"), ("Microsoft word File", "*.docx"),
                                                           ("all files", "*.*")))
        if str(save_loc).split(".")[-1] == "docx":
            document = docx.Document()
            for paragraph in text.split("\n\n"):
                document.add_paragraph(paragraph)
            document.save(save_loc)

        elif str(save_loc).split(".")[-1] == "txt":
            with open(save_loc, "w") as file:
                file.write(text)
    else:
        save_loc = f"{pathlib.Path().resolve()}/file.docx"
        print(save_loc)
        document = docx.Document()
        for paragraph in text.split("\n\n"):
            document.add_paragraph(paragraph)
        document.save(save_loc)

    create_new_doc()


def record_time(*args):
    time_now = datetime.datetime.now()
    Globals.last_press.append(time_now)

    reset_timer()
    timer_label.config(text=Globals.counter)


def timer():
    Globals.counter -= 1
    time_now = datetime.datetime.now()
    Globals.last_press.reverse()
    last_time = Globals.last_press[0]
    time_dif = (time_now - last_time).seconds
    if time_dif > 10:
        print("TIME OUT!")
        tk_message.showwarning(title="DAMN!", message="Seems like the timer ended.\nWell, too bad....")
        create_new_doc()

    timer_label.config(text=Globals.counter)
    root.after(1000, timer)


init()
root = tk.Tk()
Globals.last_press = [datetime.datetime.now()]
Globals.counter = 10

create_new_doc_button = ttk.Button(root, text="Create New Document", command=create_new_doc,
                                   width=30)
create_new_doc_button.grid(row=0, column=0, padx=(0, 5))

timer_label = ttk.Label(root, text=Globals.counter, width=30)
timer_label.configure(anchor="center")
timer_label.grid(row=0, column=1, padx=(0, 5))

save_doc_button = ttk.Button(root, text="Save Document", command=save_doc, width=30)
save_doc_button.grid(row=0, column=2)

main_entry = tk.Text(root, width=120)
main_entry.bind("<Key>", record_time)
main_entry.grid(row=1, column=0, columnspan=3, ipady=100)

root.after(1000, timer)

if __name__ == "__main__":
    root.mainloop()
